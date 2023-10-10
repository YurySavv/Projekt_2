import os
import sqlite3
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx import Document
import openpyxl
import json
from datetime import datetime

with open('path.json', 'r', encoding='utf-8') as f:
    text = json.load(f)

label_dict = {'Продажа': ['№ п/п', 'Наименование', 'Количество', 'Цена за ед.', 'Стоимость', 'Покупатель',
                          'Склад'],
              'Приемка_товара': ['№ п/п', 'Наименование', 'Количество', 'Цена за ед.', 'Стоимость', 'Поставщик',
                                 'Склад'],
              'Перемещение': ['№ п/п', 'Наименование', 'Количество', 'Цена за ед.', 'Стоимость', 'Склад приемник',
                              'Склад отправитель'],
              'Списание товара': ['№ п/п', 'Наименование', 'Количество', 'Цена за ед.', 'Стоимость', 'Причина списания',
                                  'Склад']}
type_documets = {'Продажа': 'Счет-фактура', 'Приемка_товара': 'Накладная', 'Перемещение': 'Акт перемещения',
                 'Списание товара': 'Акт списания'}
conn = sqlite3.connect('inventory_management.db')


def gen_edo(type_oper, list_data, employee_id):
    with conn:
        bd_index = conn.execute('SELECT COUNT(*) FROM Documents').fetchone()[0]

    path = text[type_oper]
    wb = openpyxl.Workbook()
    document = Document()
    if os.path.exists(path.split('\\')[0]):
        pass
    else:
        os.mkdir(path.split('\\')[0])
    if os.path.exists(path):
        pass
    else:
        os.mkdir(path)
    now = datetime.now()
    date_creation = str(now)[:10]
    wb.create_sheet(title=f'{type_oper}', index=0)
    wb.remove(wb['Sheet'])
    sheet = wb[f'{type_oper}']
    for row_index, row in enumerate((label_dict[type_oper], *list_data)):
        if row_index != 0:
            row.insert(0, row_index)
        for col_index, value in enumerate(row):
            cell = sheet.cell(row=row_index + 1, column=col_index + 1)
            cell.value = value
    #
    wb.save(f'{path}\\№_{bd_index}_{date_creation}.xlsx')
    with conn:
        conn.execute(f"""INSERT INTO Documents (type, content, employee_id, created_at)
                 VALUES (?, ?, ?, CURRENT_TIMESTAMP)""", (type_documets[type_oper],
                                                          f'№_{bd_index + 1}_{date_creation}.xlsx',
                                                          employee_id))
    os.startfile(f'{path}\\№_{bd_index}_{date_creation}.xlsx', 'open')

    styles = document.styles
    style = styles.add_style('Утверждающий', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style1 = styles.add_style('Заголовки', WD_STYLE_TYPE.PARAGRAPH)
    style1.font.name = 'Times New Roman'
    style1.font.size = Pt(16)
    style1.font.bold = True
    paragraph0 = document.add_paragraph("Утверждаю", style='Заголовки')
    paragraph0.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph1 = document.add_paragraph(
        """ Руководитель____________________
________________________________
________________________________
(дата/подпись)""", style='Утверждающий')
    paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for x in range(4):
        document.add_paragraph()
    paragraph2 = document.add_paragraph(f"{type_oper.replace('_', ' ')}", style='Заголовки')
    paragraph2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = document.add_table(rows=1 + len(list_data), cols=len(label_dict[type_oper]))
    table.style = 'Table Grid'
    head_cells = table.rows[0].cells
    for i, item in enumerate(label_dict[type_oper]):
        p = head_cells[i].paragraphs[0]
        p.add_run(item).italic = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for list1 in list_data:
        cells = table.rows[list_data.index(list1) + 1].cells
        for i, item in enumerate(list1):
            if type(item) is not str:
                item = str(item)
            p = cells[i].paragraphs[0]
            p.add_run(item)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph()
    document.add_paragraph('Отвественный ________________________\n'
                           'Подпись _____________________________', style='Утверждающий')
    with conn:
        conn.execute(f"""INSERT INTO Documents (type, content, employee_id, created_at)
         VALUES (?, ?, ?, CURRENT_TIMESTAMP)""", (type_documets[type_oper], f'№_{bd_index}_{date_creation}.docx',
                                                  employee_id))

    document.save(f'{path}\\№_{bd_index}_{date_creation}.docx')
    os.startfile(f'{path}\\№_{bd_index}_{date_creation}.docx', 'open')
    if type_oper == 'Перемещение':
        dict_moving = {}
        for item in list_data:
            move = {"id_склада-отправителя": item[5], "id_склада-получателя": item[4], "наименование": item[0],
                    "количество": item[1]}
            dict_moving.update({list_data.index(item): move})
        with open(f'{path}\\№_{bd_index}_{type_oper}_{date_creation}.json', 'w', encoding="utf-8") as moving:
            moving.write(json.dumps(dict_moving))


# generation_Word('Приемка_товара', 5)
# list_data1 = [['Телефон', 2, 4, 8, 'Рога и копыта', 15],
#               ['Стулья', 5, 4, 20, 'Рога и копыта', 16]]
# #
# gen_edo('Приемка_товара', list_data1, 5)
